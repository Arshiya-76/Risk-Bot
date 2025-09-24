# backend.py

import os
import re
import json
import docx
import io
import pypdfium2 as pdfium
import google.generativeai as genai
from dotenv import load_dotenv

load_dotenv()


def get_text_from_file(uploaded_file):
    file_bytes = uploaded_file.getvalue()
    file_type = uploaded_file.type
    try:
        if file_type == "application/pdf":
            doc = pdfium.PdfDocument(file_bytes)
            text = "".join(page.get_textpage().get_text_range() + "\n" for page in doc)
            return text
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join(para.text for para in doc.paragraphs)
            return text
        elif file_type == "text/plain":
            return file_bytes.decode('utf-8')
        else:
            return f"Error: Unsupported file type '{file_type}'."
    except Exception as e:
        return f"Error processing file: {e}"


def get_ai_analysis(raw_text: str, language: str = 'en') -> dict:
    def _segment_into_clauses(full_text):
        if not isinstance(full_text, str): return []
        pattern = r'\n\s*\d+\.\s'
        matches = re.finditer(pattern, full_text)
        start_indices = [match.start() for match in matches]
        if not start_indices: return [full_text.strip()] if len(full_text.strip()) > 20 else []
        clauses = []
        preamble = full_text[:start_indices[0]].strip()
        if len(preamble) > 20: clauses.append(preamble)
        for i in range(len(start_indices) - 1):
            start, end = start_indices[i], start_indices[i+1]
            clause_text = full_text[start:end].strip()
            if len(clause_text) > 20: clauses.append(clause_text)
        last_clause_text = full_text[start_indices[-1]:].strip()
        if len(last_clause_text) > 20: clauses.append(last_clause_text)
        return clauses

    def _get_llm_json(clauses_list, language):
        api_key = os.getenv("GOOGLE_API_KEY")
        if not api_key: return {"error": "Google API key not found."}
        genai.configure(api_key=api_key)
        full_contract_text = "\n\n".join(clauses_list)
        lang_map = {'en': 'English', 'hi': 'Hindi'}
        lang_name = lang_map.get(language, "English")
        
        system_prompt = f"""
        You are an expert AI legal assistant for Indian SMBs. Your task is to analyze a contract from an Indian SMB owner's perspective. The contract is in {lang_name}, and your analysis must also be in {lang_name}.
        You MUST provide your output in a single, valid JSON object.
        The JSON object must contain "summary_analysis" and "clause_analysis" keys.
        "summary_analysis": An object containing "contract_type", "involved_parties", "important_dates" (a list of objects, each with "date" and "context" keys), "sections_summary" (a list of objects, each with "section_name" and "simple_explanation" keys), "overall_risk_score" (1-100), "executive_summary", "key_risk_areas".
        "clause_analysis": A list of objects, each containing "risk_level" ("High", "Medium", or "Low"), "explanation", "identified_issue", "mitigation_suggestion".
        """
        
        generation_config = {"response_mime_type": "application/json", "max_output_tokens": 8192}
        model = genai.GenerativeModel('gemini-1.5-flash', system_instruction=system_prompt, generation_config=generation_config)
        try:
            response = model.generate_content(f"Please analyze the following contract text:\n\n---\n{full_contract_text}\n---")
            return json.loads(response.text)
        except Exception as e:
            return {"error": f"An error occurred during LLM analysis: {e}"}

    clauses = _segment_into_clauses(raw_text)
    return _get_llm_json(clauses, language)